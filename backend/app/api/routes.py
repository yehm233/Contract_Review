import logging
from io import BytesIO
from uuid import uuid4

from fastapi import APIRouter
from pydantic import BaseModel

from app.core.config import MINIO_ACCESS_KEY, MINIO_BUCKET, MINIO_ENDPOINT, MINIO_SECRET_KEY
from app.db.database import SessionLocal
from app.db.models import Contract, ContractStatus
from app.schemas.review import ReviewStartRequest
from app.services.review import run_review
from app.services.storage import save_onlyoffice_file, get_minio_client

logger = logging.getLogger("app.routes")

router = APIRouter()


@router.post("/api/onlyoffice/callback")
async def onlyoffice_callback(payload: dict):
    status = payload.get("status")
    url = payload.get("url")
    key = payload.get("key", str(uuid4()))
    logger.info("OnlyOffice callback | status=%s key=%s", status, key)
    if status == 2 and url:
        try:
            result = await save_onlyoffice_file(url, key)
            logger.info("OnlyOffice file saved | %s", result)
        except Exception:
            logger.exception("Failed to save OnlyOffice file | key=%s", key)
    return {"error": 0}


@router.post("/api/review/start")
async def review_start(req: ReviewStartRequest):
    logger.info("Review start | contract_id=%s", req.contract_id)
    async with SessionLocal() as session:
        # Web 提交时自动创建 contract 记录
        contract = await session.get(Contract, req.contract_id)
        if not contract:
            contract = Contract(
                contract_id=req.contract_id,
                file_name="web-submission",
                minio_url="",
                status=ContractStatus.DRAFT,
            )
            session.add(contract)
            await session.commit()
            logger.info("Auto-created contract | contract_id=%s", req.contract_id)
        return await run_review(session, req.contract_id, req.contract_text)


class GenerateRequest(BaseModel):
    party_a: str
    party_b: str
    subject: str = "合作事项"
    amount: str = ""
    duration: str = ""
    jurisdiction: str = ""
    extra: str = ""


class SaveDocRequest(BaseModel):
    contract_id: str
    contract_text: str


@router.post("/api/contract/save-doc")
async def save_contract_as_doc(req: SaveDocRequest):
    """将合同文本保存为 .docx 文件到 MinIO，返回 presigned URL 供 OnlyOffice 编辑。"""
    from docx import Document

    doc = Document()
    for line in req.contract_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
        elif len(stripped) < 50 and not stripped.endswith("。") and not stripped.endswith("；"):
            doc.add_heading(stripped, level=1 if len(stripped) < 20 else 2)
        else:
            doc.add_paragraph(stripped)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    size = buf.getbuffer().nbytes

    mc = get_minio_client()
    if not mc.bucket_exists(MINIO_BUCKET):
        mc.make_bucket(MINIO_BUCKET)

    obj_name = f"contracts/{req.contract_id}.docx"
    mc.put_object(
        MINIO_BUCKET, obj_name, buf, size,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    presigned_url = mc.presigned_get_object(MINIO_BUCKET, obj_name, expires=3600)
    logger.info("Contract doc saved | contract_id=%s obj=%s", req.contract_id, obj_name)
    return {"url": presigned_url, "object": obj_name}


@router.post("/api/contract/generate")
async def contract_generate(req: GenerateRequest):
    logger.info("Contract generate | party_a=%s party_b=%s", req.party_a, req.party_b)
    from app.services.langchain_review import LangChainReviewWorkflow, LangChainWorkflowUnavailable

    workflow = LangChainReviewWorkflow()
    tool_map = await workflow._load_tool_map()

    if "draft_contract_from_template" not in tool_map:
        return {"detail": "MCP server 不支持 draft_contract_from_template 工具"}

    form_data = {
        "party_a": req.party_a,
        "party_b": req.party_b,
        "subject": req.subject,
    }
    if req.amount:
        form_data["金额"] = req.amount
    if req.duration:
        form_data["期限"] = req.duration
    if req.jurisdiction:
        form_data["管辖"] = req.jurisdiction
    if req.extra:
        form_data["其他要求"] = req.extra

    try:
        result = await tool_map["draft_contract_from_template"].ainvoke(form_data)
        return {"contract_text": result}
    except Exception as exc:
        logger.exception("Contract generation failed")
        return {"detail": f"合同生成失败: {exc}"}
